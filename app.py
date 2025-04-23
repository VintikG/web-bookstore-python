import sqlite3
from flask import Flask, render_template, redirect, request, flash, send_from_directory
from werkzeug.exceptions import abort
import os
import pprint
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = b'my)secret)key'
UPLOAD_FOLDER = 'contracts'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/')
def index():
    return redirect("/contracts")

###############
###############
############### Контракты 
###############
###############

@app.route('/contracts')
def contracts():
    """ Страница-список - получение всех контрактов """

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, books, clients
        WHERE contracts.book_id = books.id_book and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('contracts.html', contracts=pos)


def get_contract(item_id):
    """ Получение одного контракта из БД """
    
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, books, clients
        WHERE contracts.book_id = books.id_book and contracts.client_id = clients.id_client and contracts.id_contract = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    """ Страница-карточка - 1 контракт """

    pos = get_contract(contract_id)
    return render_template('contract.html', contract=pos)


@app.route('/new_contract', methods=('GET', 'POST'))
def new_contract():
    conn = get_db_connection()
    cursor = conn.cursor()

    buyers = cursor.execute("SELECT id, name FROM buyers").fetchall()
    employees = cursor.execute("SELECT id, name FROM employees").fetchall()

    # Инициализация значений для полей формы
    number = request.form.get('number', '')
    date = request.form.get('date', '')
    service_id = request.form.get('service_id', '')
    buyer_id = request.form.get('buyer_id', '')
    finish_price = request.form.get('finish_price', '')
    employee_id = request.form.get('employee_id', '')

    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            buyer_id = int(request.form['buyer_id'])
            service_id = int(request.form['service_id'])
            finish_price = float(request.form['finish_price'])
            employee_id = int(request.form['employee_id'])
        except (ValueError, KeyError):
            flash('Некорректные данные')
            return render_template(
                'new_contract.html',
                buyers=buyers, employees=employees,
                number=number, date=date, service_id=service_id,
                buyer_id=buyer_id, finish_price=finish_price, employee_id=employee_id
            )

        if not all([number, date, buyer_id, service_id, finish_price, employee_id]):
            flash('Не все поля заполнены')
        else:
            cursor.execute(
                """INSERT INTO contracts 
                   (number, date, buyer_id, service_id, finish_price, 
                   employee_id) 
                   VALUES (?, ?, ?, ?, ?, ?)""",
                (number, date, buyer_id, service_id, finish_price, employee_id)
            )
            conn.commit()
            new_contract_id = cursor.lastrowid
            conn.close()
            return redirect(f'/contract/{new_contract_id}')

    conn.close()
    return render_template(
        'new_contract.html',
        buyers=buyers, employees=employees,
        number=number, date=date, service_id=service_id,
        buyer_id=buyer_id, finish_price=finish_price, employee_id=employee_id
    )

@app.route('/generate_contract', methods=('GET', 'POST'))
def generate_contract():
    """ Страница генерации договора """

    # переменные шаблона
    id = int(request.args.get('id_contract'))
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, books, clients, employees
            WHERE contracts.book_id = books.id_book and 
                        contracts.client_id = clients.id_client and 
                        contracts.employee_id = employees.id_employee and
                        contracts.id_contract = ?
                        """, (id,)).fetchone()
    conn.close()
    contract_params = {
            'CLIENT_BIRTHDATE': 'Дата рождения клиента',
            'CLIENT_BIRTHPLACE': 'Место рождения клиента',
            'CLIENT_REG_ADDRESS': 'Адрес регистрации',
            'CLIENT_PASSPORT_DATE': 'Дата выдачи паспорта клиента',
            'CLIENT_PASSPORT_DEPARTMENT': 'Подразделение, выдавшее паспорт клиента',
            'CLIENT_PASSPORT_DEPCODE': 'Код подразделения, выдавшего паспорт клиента',
            'BOOKS_QUANITY': 'Количество выдаваемых книг',
            'BOOKS_LIST': 'Список книг',
            'BOOKS_CONDITION': 'Состояние книг на момент выдачи',
            'RETURN_METHOD': 'Способ возврата книг',
            'CONTRACT_TERMS': 'Полные условия договора',
            'CONTRACT_VALIDATION': 'Подтверждение действительности договора',
            'CONTRACT_MODIFICATION': 'Возможность изменения условий договора',
            'CONTRACT_DEPOSIT': 'Залог за выданные книги',
            'CONTRACT_FINE': 'Штраф за просрочку договора',
            'CONTRACT_LASTDATE': 'Срок действия договора',
            'CONTRACT_CANCEL_BEFORE': 'Срок информирования намерении досрочного расторжения договора (дни)',
            'CONTRACT_CANCEL_AFTER': 'Срок расторжения договора после информирования (дни)',
            'EMPLOYEE_FULLNAME': 'ФИО сотрудника библиотеки', }
    contract_params_auto = {
            'CONTRACT_NUMBER': ['Номер договора', pos['number']],
            'CONTRACT_DATE': ['Дата подписания договора', pos['date']],
            'CLIENT_FULLNAME': ['ФИО клиента', pos['name']],
            'CLIENT_PASSPORT_NUMBER': ['Серия и номер паспорта клиента', pos['passport']],
            'EMPLOYEE_POSITION': ['Должность сотрудника библиотеки', pos['position']],}
    
    if request.method == 'POST':
        # создание нового документа
        result_params =  request.form.to_dict()
        create_contract(id, result_params)
        return redirect(f'/contract/{id}')

    # скачивание файла, если он заполнен
    filename = f"договор {pos['number']} от {pos['date']}.docx"
    if os.path.exists(os.path.join('contracts', filename)):
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)
    else:
        # отрисовка формы заполнения
        flash('Договор не сформирован, заполните его')
        return render_template('generate_contract.html', 
                               contract=pos, contract_params=contract_params, auto_params=contract_params_auto)


def create_contract(id, contract_params):
    """ Создание нового документа по шаблону """

    template = os.path.join('contracts', 'contract_template.docx')
    result = os.path.join('contracts', f"договор {contract_params['CONTRACT_NUMBER']} от {contract_params['CONTRACT_DATE']}.docx")
    
    template_doc = Document(template)
    for key, value in contract_params.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, f'=={key}==', value)
        for table in template_doc.tables:
            replace_text_in_tables(table, key, value)
    template_doc.save(result)


def replace_text(paragraph, key, value):
  """ Работа docx - заполнение параграфов """

  if key in paragraph.text:
    paragraph.text = paragraph.text.replace(key, value)


def replace_text_in_tables(table, key, value):
  """ Работа docx - заполнение таблиц """

  for row in table.rows:
    for cell in row.cells:
      if key in cell.text:
        cell.text = cell.text.replace(key, value)

###############
###############
############### Книги
###############
###############

@app.route('/books')
def books():
    """ Страница-список - получение всех книг """

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, books, clients
    WHERE contracts.book_id = books.id_book and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('books.html', books=pos)


def get_book(item_id):
    """ Получение одной книги из БД """

    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, books, clients
        WHERE contracts.book_id = books.id_book and contracts.client_id = clients.id_client and books.id_book = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/book/<int:book_id>')
def book(book_id):
    """ Страница-карточка - 1 книга """

    pos = get_book(book_id)
    return render_template('book.html', book=pos)


@app.route('/new_book', methods=('GET', 'POST'))
def new_book():
    """ Страница-добавления новой книги """

    if request.method == 'POST':
        # добавление новой книги в БД поле заполнения формы
        book_name = ", ".join([request.form['book_name'],

        ]).lstrip(", ")
        try:
            author = str(request.form['author'])
            pages = int(request.form['pages'])
            type_of_book = str(request.form['type_of_book'])
            description = request.form['description']
            client_id = int(request.form.get('owner'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 0
        if not client_id > 0:
            flash('Не все поля заполнены')
        else:
            if not (book_name and author and pages and type_of_book):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                conn.execute("INSERT INTO 'books' ('book_name', 'author', 'pages', 'type_of_book', 'description', 'client_id')  VALUES (?, ?, ?, ?, ?, ?)",
                    (book_name, author, pages, type_of_book, description, client_id))
                conn.commit()
                conn.close()
                return redirect('/new_contract')
    
    # отрисовка формы
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM clients""").fetchall()
    conn.close()
    return render_template('new_book.html', clients=pos)

#
# Клиенты 
#

@app.route('/clients')
def clients():
    """ Страница-список - получение всех клиентов """
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM clients
    """).fetchall()
    conn.close()
    return render_template('clients.html', clients=pos)


def get_client(item_id):
    """ Получение одного клиента из БД """

    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, books, clients
        WHERE contracts.client_id = clients.id_client and contracts.book_id = books.id_book and clients.id_client = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/client/<int:id_client>')
def client(id_client):
    """ Страница-карточка - 1 клиент """
    pos = get_client(id_client)
    return render_template('client.html', client=pos)

@app.route('/new_client', methods=('GET', 'POST'))
def new_client():
    """ Страница-добавления нового клиента """

    if request.method == 'POST':
        # добавление нового клиента в БД после заполнения формы
        name = ", ".join([request.form['name'],
         ]).lstrip(", ")
        try:
            phone_number =str (request.form['phone_number'])
            email = str(request.form['email'])
            passport = str(request.form['passport'])
        except ValueError:
            flash("Некорректные значения")
        else:
            if not (name and phone_number and email and passport):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                conn.execute("INSERT INTO clients (name, phone_number, email, passport) VALUES (?, ?, ?, ?)",
                             (name, phone_number, email, passport))
                conn.commit()
                conn.close()
                return redirect('/clients')

    # отрисовка формы
    return render_template('new_client.html')



#
# Отчеты 
#

@app.route('/reports')
def reports():
    abort(404)

#
# Сотрудники 
#

@app.route('/employees')
def employees():
    """ Страница-список - получение всех сотрудников """
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM employees
        """).fetchall()
    conn.close()
    return render_template('employees.html', employees=pos)

def get_employee(item_id):
    """ Получение одного сотрудника из БД """

    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, books, employees
        WHERE contracts.employee_id = employees.id_employee and contracts.book_id = books.id_book and employees.id_employee = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item

@app.route('/employee/<int:id_employee>')
def employee(id_employee):
    """ Страница-карточка - 1 сотрудник """
    pos = get_employee(id_employee)
    return render_template('employee.html', employee=pos)
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.route('/new_employee', methods=('GET', 'POST'))
def new_employee():
    """ Страница-добавления нового сотрудника """

    if request.method == 'POST':
        # добавление нового сотрудника в БД после заполнения формы
        name = ", ".join([request.form['name'],
         ]).lstrip(", ")
        try:
            phone_number =str(request.form['phone_number'])
            email = str(request.form['email'])
            position = str(request.form['position'])
            department = str(request.form['department'])
        except ValueError:
            flash("Некорректные значения")
        else:
            if not (name and phone_number and email and position and department):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                conn.execute("INSERT INTO employees (name, phone_number, email, position, department) VALUES (?, ?, ?, ?, ?)",
                             (name, phone_number, email, position, department))
                conn.commit()
                conn.close()
                return redirect('/employees')

    # отрисовка формы
    return render_template('new_employee.html')

if __name__ == '__main__':
    app.run()